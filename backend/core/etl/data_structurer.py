"""数据结构化器

负责将非结构化数据转换为结构化格式，支持多种数据源和格式。
"""

import json
import re
import uuid
from datetime import datetime
from typing import Dict, List, Any, Optional, Union, Tuple
from pathlib import Path
from dataclasses import dataclass
from enum import Enum

import pandas as pd
from pydantic import BaseModel, Field

from backend.utils.logger import get_logger
from backend.models.base import BaseModel as BaseModelClass

logger = get_logger(__name__)


class DataType(str, Enum):
    """数据类型枚举"""
    TEXT = "text"
    JSON = "json"
    CSV = "csv"
    XML = "xml"
    HTML = "html"
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    IMAGE = "image"
    AUDIO = "audio"
    VIDEO = "video"
    BINARY = "binary"


class StructureType(str, Enum):
    """结构化类型枚举"""
    DOCUMENT = "document"
    TABLE = "table"
    ENTITY = "entity"
    RELATION = "relation"
    METADATA = "metadata"
    CHUNK = "chunk"
    EMBEDDING = "embedding"


@dataclass
class StructureConfig:
    """结构化配置"""
    chunk_size: int = 1000
    chunk_overlap: int = 200
    extract_entities: bool = True
    extract_relations: bool = True
    extract_metadata: bool = True
    preserve_formatting: bool = False
    language: str = "zh"
    encoding: str = "utf-8"
    custom_rules: Dict[str, Any] = None

    def __post_init__(self):
        if self.custom_rules is None:
            self.custom_rules = {}


class StructuredData(BaseModel):
    """结构化数据模型"""
    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    source_id: str
    source_type: DataType
    structure_type: StructureType
    content: Dict[str, Any]
    metadata: Dict[str, Any] = Field(default_factory=dict)
    entities: List[Dict[str, Any]] = Field(default_factory=list)
    relations: List[Dict[str, Any]] = Field(default_factory=list)
    chunks: List[Dict[str, Any]] = Field(default_factory=list)
    quality_score: float = 0.0
    confidence: float = 0.0
    created_at: datetime = Field(default_factory=datetime.now)
    updated_at: datetime = Field(default_factory=datetime.now)


class DataStructurer:
    """数据结构化器
    
    负责将各种格式的非结构化数据转换为统一的结构化格式。
    """

    def __init__(self, config: Optional[StructureConfig] = None):
        """初始化数据结构化器
        
        Args:
            config: 结构化配置
        """
        self.config = config or StructureConfig()
        self.processors = self._initialize_processors()
        
    def _initialize_processors(self) -> Dict[DataType, callable]:
        """初始化数据处理器"""
        return {
            DataType.TEXT: self._process_text,
            DataType.JSON: self._process_json,
            DataType.CSV: self._process_csv,
            DataType.XML: self._process_xml,
            DataType.HTML: self._process_html,
            DataType.PDF: self._process_pdf,
            DataType.DOCX: self._process_docx,
            DataType.XLSX: self._process_xlsx,
        }

    async def structure_data(
        self,
        data: Union[str, bytes, Dict[str, Any]],
        data_type: DataType,
        source_id: str,
        structure_type: StructureType = StructureType.DOCUMENT,
        metadata: Optional[Dict[str, Any]] = None
    ) -> StructuredData:
        """结构化数据
        
        Args:
            data: 原始数据
            data_type: 数据类型
            source_id: 数据源ID
            structure_type: 结构化类型
            metadata: 元数据
            
        Returns:
            结构化数据对象
        """
        try:
            logger.info(f"开始结构化数据: {source_id}, 类型: {data_type}")
            
            # 获取处理器
            processor = self.processors.get(data_type)
            if not processor:
                raise ValueError(f"不支持的数据类型: {data_type}")
            
            # 处理数据
            processed_content = await processor(data)
            
            # 创建结构化数据对象
            structured_data = StructuredData(
                source_id=source_id,
                source_type=data_type,
                structure_type=structure_type,
                content=processed_content,
                metadata=metadata or {}
            )
            
            # 提取实体和关系
            if self.config.extract_entities:
                structured_data.entities = await self._extract_entities(processed_content)
            
            if self.config.extract_relations:
                structured_data.relations = await self._extract_relations(processed_content)
            
            # 分块处理
            if structure_type == StructureType.DOCUMENT:
                structured_data.chunks = await self._create_chunks(processed_content)
            
            # 计算质量分数
            structured_data.quality_score = self._calculate_quality_score(structured_data)
            structured_data.confidence = self._calculate_confidence(structured_data)
            
            logger.info(f"数据结构化完成: {source_id}, 质量分数: {structured_data.quality_score}")
            return structured_data
            
        except Exception as e:
            logger.error(f"数据结构化失败: {source_id}, 错误: {str(e)}")
            raise

    async def _process_text(self, data: str) -> Dict[str, Any]:
        """处理文本数据"""
        return {
            "text": data,
            "length": len(data),
            "lines": data.count('\n') + 1,
            "words": len(data.split()),
            "paragraphs": len([p for p in data.split('\n\n') if p.strip()])
        }

    async def _process_json(self, data: Union[str, Dict[str, Any]]) -> Dict[str, Any]:
        """处理JSON数据"""
        if isinstance(data, str):
            try:
                parsed_data = json.loads(data)
            except json.JSONDecodeError as e:
                raise ValueError(f"无效的JSON格式: {str(e)}")
        else:
            parsed_data = data
            
        return {
            "data": parsed_data,
            "keys": list(parsed_data.keys()) if isinstance(parsed_data, dict) else [],
            "size": len(str(parsed_data)),
            "depth": self._calculate_json_depth(parsed_data)
        }

    async def _process_csv(self, data: str) -> Dict[str, Any]:
        """处理CSV数据"""
        try:
            # 使用pandas读取CSV
            import io
            df = pd.read_csv(io.StringIO(data))
            
            return {
                "data": df.to_dict('records'),
                "columns": df.columns.tolist(),
                "rows": len(df),
                "shape": df.shape,
                "dtypes": df.dtypes.to_dict(),
                "summary": df.describe().to_dict() if len(df) > 0 else {}
            }
        except Exception as e:
            raise ValueError(f"CSV处理失败: {str(e)}")

    async def _process_xml(self, data: str) -> Dict[str, Any]:
        """处理XML数据"""
        try:
            import xml.etree.ElementTree as ET
            root = ET.fromstring(data)
            
            def xml_to_dict(element):
                result = {}
                if element.text and element.text.strip():
                    result['text'] = element.text.strip()
                
                for child in element:
                    child_data = xml_to_dict(child)
                    if child.tag in result:
                        if not isinstance(result[child.tag], list):
                            result[child.tag] = [result[child.tag]]
                        result[child.tag].append(child_data)
                    else:
                        result[child.tag] = child_data
                        
                result.update(element.attrib)
                return result
            
            return {
                "data": xml_to_dict(root),
                "root_tag": root.tag,
                "attributes": root.attrib,
                "elements_count": len(list(root.iter()))
            }
        except Exception as e:
            raise ValueError(f"XML处理失败: {str(e)}")

    async def _process_html(self, data: str) -> Dict[str, Any]:
        """处理HTML数据"""
        try:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(data, 'html.parser')
            
            # 提取文本内容
            text_content = soup.get_text(separator='\n', strip=True)
            
            # 提取结构化信息
            headings = [h.get_text(strip=True) for h in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])]
            links = [{'text': a.get_text(strip=True), 'href': a.get('href')} for a in soup.find_all('a', href=True)]
            images = [{'alt': img.get('alt', ''), 'src': img.get('src')} for img in soup.find_all('img')]
            
            return {
                "text": text_content,
                "title": soup.title.string if soup.title else "",
                "headings": headings,
                "links": links,
                "images": images,
                "tables": len(soup.find_all('table')),
                "forms": len(soup.find_all('form'))
            }
        except Exception as e:
            raise ValueError(f"HTML处理失败: {str(e)}")

    async def _process_pdf(self, data: bytes) -> Dict[str, Any]:
        """处理PDF数据"""
        # 这里应该集成OCR服务来处理PDF
        # 暂时返回基本信息
        return {
            "type": "pdf",
            "size": len(data),
            "text": "PDF内容需要OCR处理",
            "pages": 0  # 需要PDF库来获取页数
        }

    async def _process_docx(self, data: bytes) -> Dict[str, Any]:
        """处理DOCX数据"""
        try:
            from docx import Document
            import io
            
            doc = Document(io.BytesIO(data))
            
            # 提取文本内容
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text_content = '\n'.join(paragraphs)
            
            # 提取表格
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            return {
                "text": text_content,
                "paragraphs": paragraphs,
                "tables": tables,
                "paragraphs_count": len(paragraphs),
                "tables_count": len(tables)
            }
        except Exception as e:
            raise ValueError(f"DOCX处理失败: {str(e)}")

    async def _process_xlsx(self, data: bytes) -> Dict[str, Any]:
        """处理XLSX数据"""
        try:
            import io
            df = pd.read_excel(io.BytesIO(data), sheet_name=None)
            
            result = {
                "sheets": {},
                "sheet_names": list(df.keys()),
                "total_rows": 0,
                "total_columns": 0
            }
            
            for sheet_name, sheet_df in df.items():
                result["sheets"][sheet_name] = {
                    "data": sheet_df.to_dict('records'),
                    "columns": sheet_df.columns.tolist(),
                    "rows": len(sheet_df),
                    "shape": sheet_df.shape
                }
                result["total_rows"] += len(sheet_df)
                result["total_columns"] += len(sheet_df.columns)
            
            return result
        except Exception as e:
            raise ValueError(f"XLSX处理失败: {str(e)}")

    async def _extract_entities(self, content: Dict[str, Any]) -> List[Dict[str, Any]]:
        """提取实体"""
        entities = []
        
        # 简单的实体提取逻辑
        text = self._extract_text_from_content(content)
        if text:
            # 提取邮箱
            emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
            for email in emails:
                entities.append({
                    "type": "email",
                    "value": email,
                    "confidence": 0.9
                })
            
            # 提取电话号码
            phones = re.findall(r'\b(?:\+?86)?\s*1[3-9]\d{9}\b', text)
            for phone in phones:
                entities.append({
                    "type": "phone",
                    "value": phone,
                    "confidence": 0.8
                })
            
            # 提取日期
            dates = re.findall(r'\b\d{4}[-/]\d{1,2}[-/]\d{1,2}\b', text)
            for date in dates:
                entities.append({
                    "type": "date",
                    "value": date,
                    "confidence": 0.7
                })
        
        return entities

    async def _extract_relations(self, content: Dict[str, Any]) -> List[Dict[str, Any]]:
        """提取关系"""
        # 简单的关系提取逻辑
        # 实际应用中应该使用更复杂的NLP模型
        return []

    async def _create_chunks(self, content: Dict[str, Any]) -> List[Dict[str, Any]]:
        """创建文本块"""
        chunks = []
        text = self._extract_text_from_content(content)
        
        if text:
            # 简单的分块逻辑
            chunk_size = self.config.chunk_size
            overlap = self.config.chunk_overlap
            
            for i in range(0, len(text), chunk_size - overlap):
                chunk_text = text[i:i + chunk_size]
                if chunk_text.strip():
                    chunks.append({
                        "id": str(uuid.uuid4()),
                        "text": chunk_text,
                        "start_index": i,
                        "end_index": min(i + chunk_size, len(text)),
                        "length": len(chunk_text)
                    })
        
        return chunks

    def _extract_text_from_content(self, content: Dict[str, Any]) -> str:
        """从内容中提取文本"""
        if "text" in content:
            return content["text"]
        elif "data" in content and isinstance(content["data"], str):
            return content["data"]
        elif "paragraphs" in content:
            return '\n'.join(content["paragraphs"])
        else:
            return str(content)

    def _calculate_json_depth(self, obj: Any, depth: int = 0) -> int:
        """计算JSON深度"""
        if isinstance(obj, dict):
            return max([self._calculate_json_depth(v, depth + 1) for v in obj.values()], default=depth)
        elif isinstance(obj, list):
            return max([self._calculate_json_depth(item, depth + 1) for item in obj], default=depth)
        else:
            return depth

    def _calculate_quality_score(self, structured_data: StructuredData) -> float:
        """计算数据质量分数"""
        score = 0.0
        
        # 基于内容完整性
        if structured_data.content:
            score += 0.3
        
        # 基于实体数量
        if structured_data.entities:
            score += min(len(structured_data.entities) * 0.1, 0.3)
        
        # 基于关系数量
        if structured_data.relations:
            score += min(len(structured_data.relations) * 0.1, 0.2)
        
        # 基于元数据完整性
        if structured_data.metadata:
            score += min(len(structured_data.metadata) * 0.05, 0.2)
        
        return min(score, 1.0)

    def _calculate_confidence(self, structured_data: StructuredData) -> float:
        """计算置信度"""
        # 简单的置信度计算
        confidence = 0.5  # 基础置信度
        
        # 基于数据类型的置信度调整
        if structured_data.source_type in [DataType.JSON, DataType.CSV]:
            confidence += 0.3
        elif structured_data.source_type in [DataType.TEXT, DataType.HTML]:
            confidence += 0.2
        
        # 基于质量分数的调整
        confidence += structured_data.quality_score * 0.2
        
        return min(confidence, 1.0)

    async def batch_structure(
        self,
        data_items: List[Tuple[Any, DataType, str]],
        structure_type: StructureType = StructureType.DOCUMENT
    ) -> List[StructuredData]:
        """批量结构化数据
        
        Args:
            data_items: 数据项列表，每项包含(数据, 数据类型, 源ID)
            structure_type: 结构化类型
            
        Returns:
            结构化数据列表
        """
        results = []
        
        for data, data_type, source_id in data_items:
            try:
                structured = await self.structure_data(
                    data=data,
                    data_type=data_type,
                    source_id=source_id,
                    structure_type=structure_type
                )
                results.append(structured)
            except Exception as e:
                logger.error(f"批量处理失败: {source_id}, 错误: {str(e)}")
                # 创建错误记录
                error_data = StructuredData(
                    source_id=source_id,
                    source_type=data_type,
                    structure_type=structure_type,
                    content={"error": str(e)},
                    quality_score=0.0,
                    confidence=0.0
                )
                results.append(error_data)
        
        return results

    def get_supported_types(self) -> List[DataType]:
        """获取支持的数据类型"""
        return list(self.processors.keys())

    def update_config(self, config: StructureConfig) -> None:
        """更新配置"""
        self.config = config
        logger.info("数据结构化器配置已更新")